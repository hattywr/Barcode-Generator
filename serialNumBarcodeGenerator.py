import random
import string
import barcode
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime

#Generate Serial Numbers
def generate_serial(length=20):
    characters = string.ascii_uppercase + string.digits
    serial = ''.join(random.choice(characters) for _ in range(length))
    return serial

# Delete Existing Barcodes
def cleanup_barcode_files(folder="barcodes"):
    if os.path.exists(folder):
        for filename in os.listdir(folder):
            if filename.endswith(".png"):
                file_path = os.path.join(folder, filename)
                try:
                    os.remove(file_path)
                    print(f"Deleted: {filename}")
                except Exception as e:
                    print(f"Error deleting {filename}: {e}")
    else:
        print(f"Folder '{folder}' does not exist")

#generate barcode images using serial #'s
def generate_barcode(serial_number, folder="barcodes"):
    # Create the folder if it doesn't exist
    if not os.path.exists(folder):
        os.makedirs(folder)
        
    barcode_class = barcode.get_barcode_class('code128')
    if barcode_class is None:
        raise ValueError("Invalid barcode type")
    
    barcode_image_path = os.path.join(folder, f"{serial_number}_barcode.png")
    my_barcode = barcode_class(serial_number, writer=ImageWriter())
    no_png_path = barcode_image_path.replace(".png", "")
    my_barcode.save(no_png_path)
    
    return barcode_image_path

#generate desired amt of barcodes
def generate_multiple_barcodes(num_barcodes):
    barcode_data = []
    for _ in range(num_barcodes):
        serial_number = generate_serial()
        barcode_image_path = generate_barcode(serial_number)
        barcode_data.append((serial_number, barcode_image_path))
    
    return barcode_data

#Put the barcodes into a word doc - if append = true, takc them on the end, if false, overwrite the doc (Wipe it)
def insert_barcodes_into_doc(barcode_data, doc_path, append=False):
    if append and os.path.exists(doc_path):
        # Load existing document if appending
        doc = Document(doc_path)
        # Add a page break before new content

        # Add timestamp header
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_heading(f"Newly Generated At {timestamp}", level=1)
        doc.add_paragraph()  # Add some space after the header
    else:
        # Create a new document if not appending or if file doesn't exist
        doc = Document()
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_heading(f"Newly Generated At {timestamp}", level=1)
        doc.add_paragraph()  # Add some space after the header

    # Loop through the barcode data and add them to the document
    for serial_number, image_path in barcode_data:
        doc.add_paragraph(f"Serial Number: {serial_number}")
        doc.add_picture(image_path, width=Inches(4))
        doc.add_paragraph()  # Add space between each barcode
    
    # Save the document
    doc.save(doc_path)
    print(f"All barcodes {'added to' if append else 'saved in'} {doc_path}")

if __name__ == "__main__":
    # Clean up existing barcode files
    cleanup_barcode_files()
    
    # Generate and insert barcodes
    num_barcodes = 20  # Number of barcodes to generate
    append_mode = False  # Set to False to overwrite the document instead of appending
    
    barcode_data = generate_multiple_barcodes(num_barcodes)
    insert_barcodes_into_doc(barcode_data, doc_path="barcodes.docx", append=append_mode)