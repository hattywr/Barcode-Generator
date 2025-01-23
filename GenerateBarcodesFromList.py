import barcode
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime

def cleanup_barcode_files(folder="barcodes"):
    if os.path.exists(folder):
        for filename in os.listdir(folder):
            if filename.endswith(".png"):
                os.remove(os.path.join(folder, filename))

def generate_barcode(value, folder="barcodes"):
    if not os.path.exists(folder):
        os.makedirs(folder)
    
    writer = ImageWriter()
    writer.module_width = 0.4
    writer.module_height = 6.0
    
    barcode_class = barcode.get_barcode_class('code128')
    barcode_image_path = os.path.join(folder, f"{value}_barcode.png")
    my_barcode = barcode_class(str(value), writer=writer)
    my_barcode.save(barcode_image_path.replace(".png", ""))
    
    return barcode_image_path

def generate_multiple_barcodes(values):
    return [(value, generate_barcode(value)) for value in values]

def insert_barcodes_into_doc(barcode_data, doc_path, append=False):
    doc = Document(doc_path) if append and os.path.exists(doc_path) else Document()
    
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_heading(f"Generated At {timestamp}", level=1)
    doc.add_paragraph()

    for value, image_path in barcode_data:
        doc.add_paragraph(f"Value: {value}")
        doc.add_picture(image_path, width=Inches(7),height=Inches(3))
        doc.add_paragraph()
    
    doc.save(doc_path)

if __name__ == "__main__":
    cleanup_barcode_files()
    
    # Example usage with list of values
    values = ["50010101-A", "50010101-B", "50010101-C", "50010102-A", "50010102-B", "50010102-C", "50010103-A", "50010103-B", "50010103-C","50010101-D", "50010101-E", "50010101-F", "50010102-D", "50010102-E", "50010102-F", "50010103-D", "50010103-E", "50010103-F", "50010201-A", "50010201-B", "50010201-C", "50010201-D", "50010201-E", "50010201-F", "50010202-A", "50010202-B", "50010202-C", "50010202-D", "50010202-E", "50010202-F", "50010203-A", "50010203-B", "50010203-C", "50010203-D", "50010203-E", "50010203-F", "50010301-A", "50010301-B", "50010301-C", "50010301-D", "50010301-E", "50010301-F", "50010302-A", "50010302-B", "50010302-C", "50010302-D", "50010302-E", "50010302-F", "50010303-A", "50010303-B", "50010303-C", "50010303-D", "50010303-E", "50010303-F"]  # Add your values here
    barcode_data = generate_multiple_barcodes(values)
    insert_barcodes_into_doc(barcode_data, "barcodes.docx", append=False)